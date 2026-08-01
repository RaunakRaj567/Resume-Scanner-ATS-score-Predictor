import re
from pathlib import Path
from typing import Dict, Any, List
from app.utils.text_cleaning import clean_text
from app.utils.logging import get_logger

logger = get_logger(__name__)

class ResumeExtractor:
    def __init__(self):
        pass

    def extract_from_file(self, file_path: Path) -> Dict[str, Any]:
        """Extract text, section breakdown, and metadata from PDF, DOCX, or TXT file."""
        extension = file_path.suffix.lower()
        raw_text = ""

        if extension == ".pdf":
            raw_text = self._extract_pdf(file_path)
        elif extension in (".docx", ".doc"):
            raw_text = self._extract_docx(file_path)
        elif extension in (".txt", ".md"):
            raw_text = self._extract_txt(file_path)
        else:
            raise ValueError(f"Unsupported file format: {extension}")

        cleaned_text = clean_text(raw_text)
        
        # OCR Fallback if text extraction yields minimal content
        if len(cleaned_text.strip()) < 30 and extension == ".pdf":
            logger.info(f"PDF text content very low ({len(cleaned_text)} chars), attempting OCR fallback if available...")
            ocr_text = self._extract_ocr(file_path)
            if len(ocr_text) > len(cleaned_text):
                cleaned_text = clean_text(ocr_text)

        sections = self._parse_sections(cleaned_text)
        return {
            "extracted_text": cleaned_text,
            "sections": sections,
            "char_count": len(cleaned_text),
            "word_count": len(cleaned_text.split())
        }

    def _extract_pdf(self, file_path: Path) -> str:
        text = ""
        # Try pdfplumber first
        try:
            import pdfplumber
            with pdfplumber.open(file_path) as pdf:
                pages_text = [page.extract_text() or "" for page in pdf.pages]
                text = "\n".join(pages_text)
        except Exception as e:
            logger.warning(f"pdfplumber extraction failed for {file_path}: {e}, falling back to pypdf")
            try:
                from pypdf import PdfReader
                reader = PdfReader(file_path)
                pages_text = [page.extract_text() or "" for page in reader.pages]
                text = "\n".join(pages_text)
            except Exception as e2:
                logger.error(f"pypdf extraction failed as well: {e2}")
        return text

    def _extract_docx(self, file_path: Path) -> str:
        try:
            import docx
            doc = docx.Document(file_path)
            full_text = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        full_text.append(row_text)
            return "\n".join(full_text)
        except Exception as e:
            logger.error(f"DOCX extraction failed for {file_path}: {e}")
            return ""

    def _extract_txt(self, file_path: Path) -> str:
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()
        except Exception as e:
            logger.error(f"TXT extraction failed for {file_path}: {e}")
            return ""

    def _extract_ocr(self, file_path: Path) -> str:
        """Optional Tesseract OCR fallback for scanned image PDFs."""
        try:
            import pytesseract
            from pypdf import PdfReader
            # Basic OCR attempt using pdf2image if available
            try:
                from pdf2image import convert_from_path
                images = convert_from_path(file_path)
                ocr_pages = [pytesseract.image_to_string(img) for img in images]
                return "\n".join(ocr_pages)
            except Exception:
                logger.info("pdf2image not available for OCR fallback")
                return ""
        except ImportError:
            logger.info("pytesseract or PIL not installed for OCR fallback")
            return ""

    def _parse_sections(self, text: str) -> Dict[str, List[str]]:
        """Identify standard resume sections (Experience, Education, Skills, Projects, Summary)."""
        section_headers = {
            "summary": [r"^summary\b", r"^professional summary\b", r"^executive summary\b", r"^profile\b", r"^objective\b", r"^about me\b"],
            "experience": [r"^experience\b", r"^work experience\b", r"^employment\b", r"^work history\b", r"^professional experience\b", r"^relevant experience\b"],
            "education": [r"^education\b", r"^academic background\b", r"^academic qualifications\b", r"^education & certifications\b"],
            "skills": [r"^skills\b", r"^technical skills\b", r"^technologies\b", r"^core competencies\b", r"^skills & abilities\b", r"^skills & tools\b"],
            "projects": [r"^projects?\b", r"^key projects?\b", r"^personal projects?\b", r"^academic projects?\b", r"^featured projects?\b", r"^technical projects?\b", r"^project portfolio\b"]
        }

        lines = text.split("\n")
        sections = {"summary": [], "experience": [], "education": [], "skills": [], "projects": [], "other": []}
        current_section = "other"

        for line in lines:
            line_str = line.strip()
            if not line_str:
                continue

            # Check if line matches a section header (headers are usually short <= 45 chars)
            matched_header = None
            if len(line_str) <= 45:
                clean_line_lower = line_str.lower().rstrip(":#- ")
                for sec_name, patterns in section_headers.items():
                    for pat in patterns:
                        if re.search(pat, clean_line_lower):
                            matched_header = sec_name
                            break
                    if matched_header:
                        break

            if matched_header:
                current_section = matched_header
            else:
                sections[current_section].append(line_str)

        return sections

extractor = ResumeExtractor()
